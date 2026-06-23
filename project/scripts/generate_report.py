from __future__ import annotations

import json
import re
from datetime import datetime
from pathlib import Path

PROJECT_ROOT = Path(__file__).resolve().parents[1]
DOCS_DIR = PROJECT_ROOT / "docs"
REPORT_PARTS_DIR = DOCS_DIR / "report_parts"
DATABASE_DIR = PROJECT_ROOT / "database"
DIAGRAMS_DIR = PROJECT_ROOT / "diagrams"
SCREENSHOTS_DIR = PROJECT_ROOT / "screenshots"
REPORT_MD = DOCS_DIR / "重庆邮电大学访客管理系统设计报告.md"
REPORT_DOCX = DOCS_DIR / "重庆邮电大学访客管理系统设计报告.docx"

SYSTEM_NAME = "重庆邮电大学智慧访客预约与出入校管理系统"
ENGLISH_NAME = "Smart Visitor Reservation and Campus Access Management System of Chongqing University of Posts and Telecommunications"


def read_text(path: Path, default: str = "") -> str:
    if not path.exists():
        return default
    return path.read_text(encoding="utf-8", errors="replace").strip()


def strip_first_heading(text: str) -> str:
    lines = text.splitlines()
    if lines and lines[0].lstrip().startswith("#"):
        return "\n".join(lines[1:]).strip()
    return text.strip()


def fenced_file(path: Path, lang: str = "") -> str:
    text = read_text(path)
    if not text:
        return f"> 缺少文件：`{path.relative_to(PROJECT_ROOT)}`\n"
    return f"```{lang}\n{text}\n```\n"


def mermaid_from(path: Path) -> str:
    return fenced_file(path, "mermaid")


def load_manifest() -> list[dict]:
    path = SCREENSHOTS_DIR / "manifest.json"
    if not path.exists():
        return []
    return json.loads(path.read_text(encoding="utf-8"))


def parse_create_tables() -> list[str]:
    schema = read_text(DATABASE_DIR / "schema.sql")
    return re.findall(r"CREATE\s+TABLE\s+`?([a-zA-Z0-9_]+)`?", schema, flags=re.I)


def screenshot_section() -> str:
    manifest = load_manifest()
    if not manifest:
        return "> 尚未生成 `screenshots/manifest.json`，请先执行 `node scripts/capture_screenshots.js`。\n"
    lines: list[str] = []
    success_count = sum(1 for item in manifest if item.get("status") == "SUCCESS")
    lines.append(f"自动截图脚本共登记 {len(manifest)} 个页面，其中成功截图 {success_count} 个。截图文件采用固定命名，便于课程设计报告稳定引用。")
    lines.append("")
    for index, item in enumerate(manifest, start=1):
        page_name = item.get("pageName", item.get("screenshotCode", f"截图{index}"))
        file_path = item.get("filePath", "")
        desc = item.get("description", "")
        report_path = "../" + file_path.replace("\\", "/")
        lines.append(f"![{page_name}]({report_path})")
        lines.append("")
        lines.append(f"图 5-{index} {page_name}。{desc}")
        lines.append("")
    return "\n".join(lines).strip()


def sql_summary_section() -> str:
    tables = parse_create_tables()
    lines = [
        "数据库脚本位于 `database/` 目录，包含建库脚本、建表脚本、初始化数据和典型查询语句。",
        "",
        f"`schema.sql` 共定义 {len(tables)} 张业务表：`" + "`、`".join(tables) + "`。" if tables else "未从 `schema.sql` 中解析到建表语句。",
        "",
        "### 建表脚本摘要",
        "",
        "完整建表语句见 `database/schema.sql`，以下保留核心脚本，体现主键、外键、索引、字段注释和表注释设计。",
        "",
        fenced_file(DATABASE_DIR / "schema.sql", "sql"),
        "### 典型查询脚本",
        "",
        "以下 SQL 来自 `database/query_examples.sql`，覆盖历史预约、待确认、待审批、在校访客、超时未离校、部门排行、校门通行、黑名单、审批记录、车辆随行人员、操作日志、趋势和审批通过率等查询。",
        "",
        fenced_file(DATABASE_DIR / "query_examples.sql", "sql"),
    ]
    return "\n".join(lines)


def diagram_block(title: str, path: Path, description: str) -> str:
    return f"### {title}\n\n{mermaid_from(path)}\n{description}\n"


def build_markdown() -> str:
    now = datetime.now().strftime("%Y-%m-%d %H:%M:%S")
    parts = {
        "system": strip_first_heading(read_text(REPORT_PARTS_DIR / "01_系统定义.md")),
        "requirements": strip_first_heading(read_text(REPORT_PARTS_DIR / "02_需求分析.md")),
        "data_flow": strip_first_heading(read_text(REPORT_PARTS_DIR / "03_数据流程图.md")),
        "dictionary": strip_first_heading(read_text(REPORT_PARTS_DIR / "04_数据字典.md")),
        "concept": strip_first_heading(read_text(REPORT_PARTS_DIR / "05_概念结构设计.md")),
        "er": strip_first_heading(read_text(REPORT_PARTS_DIR / "06_ER图.md")),
        "logic": strip_first_heading(read_text(REPORT_PARTS_DIR / "07_逻辑结构设计.md")),
        "sql_doc": strip_first_heading(read_text(REPORT_PARTS_DIR / "08_SQL查询语句和说明.md")),
        "relation": strip_first_heading(read_text(DOCS_DIR / "关系模式.md")),
        "diagrams": strip_first_heading(read_text(DOCS_DIR / "diagrams.md")),
        "api": strip_first_heading(read_text(DOCS_DIR / "API接口说明.md")),
        "test": strip_first_heading(read_text(DOCS_DIR / "系统测试说明.md")),
        "cases": strip_first_heading(read_text(DOCS_DIR / "测试用例.md")),
        "accounts": strip_first_heading(read_text(DOCS_DIR / "测试账号说明.md")),
    }

    md: list[str] = []
    md.append(f"# {SYSTEM_NAME}设计报告")
    md.append("")
    md.append(f"英文名称：{ENGLISH_NAME}")
    md.append("")
    md.append(f"生成时间：{now}")
    md.append("")
    md.append("技术栈：Vue 3 + Vite + Element Plus + Axios + Vue Router + ECharts；Spring Boot 3 + MyBatis Plus + MySQL + JWT；MySQL 8；Playwright；Markdown + python-docx。")
    md.append("")

    md.append("## 一、系统定义")
    md.append(parts["system"])
    md.append("")

    md.append("## 二、需求分析")
    md.append("### 1. 综合需求")
    md.append("系统面向重庆邮电大学校外访客预约、校内被访人确认、部门审批、门岗核验、入校离校登记和校级统计监管等场景，要求实现从预约申请到访问记录归档的闭环管理。")
    md.append("")
    md.append("### 2. 功能需求")
    md.append(parts["requirements"])
    md.append("")
    md.append("### 3. 性能需求")
    md.append("系统应支持常规课程设计演示规模下的多角色并发访问，核心查询通过主键、外键和状态字段建立索引；统计报表以聚合查询为主，保证首页和报表页能够快速展示。")
    md.append("")
    md.append("### 4. 运行需求")
    md.append("系统运行需要 MySQL 8、JDK 17、Maven、Node.js/npm 和现代浏览器。自动截图需要 Playwright 或系统 Chrome。")
    md.append("")
    md.append("### 5. 数据流程图")
    md.append(parts["data_flow"])
    md.append("")
    md.append(mermaid_from(DIAGRAMS_DIR / "data_flow.mmd"))
    md.append("")
    md.append("### 6. 数据字典")
    md.append(parts["dictionary"])
    md.append("")

    md.append("## 三、系统设计")
    md.append("### 1. 概念结构设计")
    md.append(parts["concept"])
    md.append("")
    md.append("### 2. E-R 图")
    md.append(parts["er"])
    md.append("")
    md.append(mermaid_from(DIAGRAMS_DIR / "er_diagram.mmd"))
    md.append("")
    md.append("### 3. 逻辑结构设计")
    md.append(parts["logic"])
    md.append("")
    md.append("### 4. 关系模式")
    md.append(parts["relation"])
    md.append("")
    md.append("### 5. 系统功能模块图")
    md.append(diagram_block("系统功能模块图", DIAGRAMS_DIR / "system_module.mmd", "系统功能模块图展示访客、被访人、审批人员、门岗安保、系统管理员和校级管理人员对应的功能边界。"))
    md.append("### 6. 其它设计图")
    md.append(diagram_block("用户用例图", DIAGRAMS_DIR / "use_case.mmd", "用例图说明不同角色可触达的业务能力和权限范围。"))
    md.append(diagram_block("数据库表关系图", DIAGRAMS_DIR / "table_relation.mmd", "数据库表关系图体现用户权限、访客预约、审批、通行凭证、出入校记录和日志等表之间的关联。"))
    md.append(diagram_block("系统架构图", DIAGRAMS_DIR / "system_architecture.mmd", "系统架构图说明前端、后端、数据库、截图脚本和报告生成脚本之间的协作关系。"))
    md.append(diagram_block("访客预约审批流程图", DIAGRAMS_DIR / "visitor_workflow.mmd", "访客预约审批流程图对应核心业务链路：提交申请、黑名单检查、被访人确认、部门审批和通行凭证生成。"))
    md.append(diagram_block("门岗核验入校流程图", DIAGRAMS_DIR / "gate_check_workflow.mmd", "门岗核验入校流程图说明门岗按通行凭证核验、登记入校、登记离校和处理异常状态的过程。"))
    md.append(diagram_block("自动截图和报告生成流程图", DIAGRAMS_DIR / "automation_report_workflow.mmd", "自动化流程图说明系统运行截图、manifest 生成、报告素材读取和 Word 导出的过程。"))
    md.append("")

    md.append("## 四、详细设计")
    md.append("### 1. SQL 查询语句和说明")
    md.append(parts["sql_doc"])
    md.append("")
    md.append(sql_summary_section())
    md.append("")

    md.append("## 五、系统实现与测试")
    md.append("### 1. 开发平台和工具选择")
    md.append("后端采用 Spring Boot 3、MyBatis Plus、MySQL Connector/J、Spring Security 和 JWT；前端采用 Vue 3、Vite、Element Plus、Axios、Vue Router 和 ECharts；数据库采用 MySQL 8；自动截图采用 Playwright；报告自动生成采用 Python 和 python-docx。")
    md.append("")
    md.append("### 接口实现摘要")
    md.append(parts["api"][:12000])
    md.append("")
    md.append("### 测试账号")
    md.append(parts["accounts"])
    md.append("")
    md.append("### 2. 系统测试")
    md.append(parts["test"])
    md.append("")
    md.append("### 测试用例表")
    md.append(parts["cases"])
    md.append("")
    md.append("### 3. 代表性运行界面")
    md.append(screenshot_section())
    md.append("")

    md.append("## 六、总结")
    md.append("本系统围绕重庆邮电大学访客管理场景，完成了从需求分析、概念结构设计、逻辑结构设计、MySQL 建库建表、后端接口、角色权限、访客预约审批流程、门岗核验、统计报表、自动截图到课程设计报告自动生成的完整闭环。数据库设计覆盖用户、角色、权限、部门、校门、访客、车辆、随行人员、预约申请、审批记录、通行凭证、出入校记录、黑名单、通知、日志、截图和报告等核心数据对象，能够支撑课程设计要求中的数据流程图、数据字典、E-R 图、关系模式、SQL 查询和系统测试内容。")
    md.append("")
    md.append("通过本次课程设计，系统不仅实现了基础 CRUD，还体现了数据库课程中实体联系分析、关系模式转换、约束设计、索引设计、典型查询设计和数据一致性维护等重点内容。后续可继续增强接口自动化测试、细化审计日志、优化前端分包加载，并在真实部署环境中补充 HTTPS、网关限流和更严格的数据脱敏策略。")
    md.append("")
    return "\n".join(md).replace("\r\n", "\n")


def add_code_block(doc, text: str):
    paragraph = doc.add_paragraph()
    run = paragraph.add_run(text[:6000])
    run.font.name = "Consolas"
    run.font.size = docx.shared.Pt(8)


def export_docx(markdown: str) -> str:
    try:
        global docx
        import docx
        from docx.shared import Inches, Pt
    except Exception as exc:
        note = REPORT_DOCX.with_suffix(".docx.txt")
        note.write_text(f"未能导出 Word：缺少 python-docx。错误：{exc}\n请执行 python -m pip install python-docx 后重试。", encoding="utf-8")
        return str(note)

    document = docx.Document()
    document.core_properties.title = f"{SYSTEM_NAME}设计报告"
    document.core_properties.subject = "数据库原理课程设计"
    document.add_heading(f"{SYSTEM_NAME}设计报告", level=0)
    document.add_paragraph(f"英文名称：{ENGLISH_NAME}")

    in_code = False
    code_lines: list[str] = []
    for raw_line in markdown.splitlines():
        line = raw_line.rstrip()
        if line.startswith("```"):
            if in_code:
                add_code_block(document, "\n".join(code_lines))
                code_lines = []
                in_code = False
            else:
                in_code = True
            continue
        if in_code:
            code_lines.append(line)
            continue
        image_match = re.match(r"!\[(.*?)\]\((.*?)\)", line)
        if image_match:
            alt, image_path = image_match.groups()
            resolved = (DOCS_DIR / image_path).resolve()
            if resolved.exists():
                document.add_picture(str(resolved), width=Inches(5.8))
                document.add_paragraph(alt)
            else:
                document.add_paragraph(f"[缺少图片] {alt}: {image_path}")
            continue
        if not line.strip():
            continue
        if line.startswith("# "):
            continue
        if line.startswith("## "):
            document.add_heading(line[3:].strip(), level=1)
        elif line.startswith("### "):
            document.add_heading(line[4:].strip(), level=2)
        elif line.startswith("#### "):
            document.add_heading(line[5:].strip(), level=3)
        else:
            document.add_paragraph(line)
    if code_lines:
        add_code_block(document, "\n".join(code_lines))
    document.save(REPORT_DOCX)
    return str(REPORT_DOCX)


def main() -> None:
    DOCS_DIR.mkdir(parents=True, exist_ok=True)
    markdown = build_markdown()
    REPORT_MD.write_text(markdown, encoding="utf-8")
    docx_path = export_docx(markdown)
    print(f"Markdown report generated: {REPORT_MD}")
    print(f"Word report generated: {docx_path}")


if __name__ == "__main__":
    main()